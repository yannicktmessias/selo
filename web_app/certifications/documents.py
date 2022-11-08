from .models import Certification, Page, EvaluationReport
import docx
import os
from django.conf import settings

templates_path = os.path.join(settings.BASE_DIR, 'templates')

def evaluation_info_dictionary(sei_number, evaluation_date, evaluator):
    certification = Certification.objects.get(sei_number=sei_number)
    evaluation_info = {
        '<sei_number>': f'{sei_number[0:4]}.{sei_number[4:8]}/{sei_number[8:15]}-{sei_number[15]}',
        '<certification_code>': certification.code,
        '<applicant>': certification.applicant.name,
        '<domain>': certification.domain,
        '<methodology>': 'e-MAG e WCAG 2.0',
        '<evaluation_date>': '/'.join(reversed(str(evaluation_date).split('-'))),
        '<evaluator>': evaluator,
    }
    return evaluation_info

def page_evaluations_dictionary(sei_number, evaluation_date):
    certification = Certification.objects.get(sei_number=sei_number)
    pages = Page.objects.filter(certification=certification)
    page_evaluations = []
    for i, page in enumerate(pages):
        page_reports = EvaluationReport.objects.filter(page=page)
        date_page_report = page_reports.filter(
            creation_date_time__year=evaluation_date.year,
            creation_date_time__month=evaluation_date.month,
            creation_date_time__day=evaluation_date.day,
        )
        if date_page_report:
            grade = date_page_report[0].grade
        else:
            grade = 0
        page_evaluations.append([
            str(i),
            page.url,
            str(grade) + '%',
        ])
    return page_evaluations

def make_list_of_pages_document(evaluation_info, page_evaluations):
    document = docx.Document(os.path.join(templates_path, 'documents/list_of_pages.docx'))

    for paragraph in document.paragraphs:
        if not paragraph == '':
            for run in paragraph.runs:
                for label in evaluation_info.keys():
                    if label in run.text:
                        run.text = run.text.replace(label, evaluation_info[label])
                        
    '''
    if evaluation_info['<evaluator2>'].strip() != '' or evaluation_info['<evaluator3>'].strip() != '':
        cell = document.tables[0].rows[-1].cells[0]
        cell.paragraphs[0].runs[0].text = 'Avaliadores:'

    if evaluation_info['<evaluator2>'].strip() != '':
        document.tables[0].add_row()
        cell = document.tables[0].rows[-1].cells[0]
        cell.merge(document.tables[0].rows[-2].cells[0])

        cell = document.tables[0].rows[-1].cells[1]
        
        cell.vertical_alignment = document.tables[0].rows[-2].cells[1].vertical_alignment
        base_paragraph = document.tables[0].rows[-2].cells[1].paragraphs[0]
        cell.paragraphs[0].alignment = base_paragraph.alignment

        run = cell.paragraphs[0].add_run()
        run.text = evaluation_info['<evaluator2>']

        base_font = document.tables[0].rows[-2].cells[1].paragraphs[0].runs[0].font
        run.font.name = base_font.name
        run.font.size = base_font.size
        run.font.bold = base_font.bold
        run.font.underline = base_font.underline
        run.font.color.rgb = base_font.color.rgb

    if evaluation_info['<evaluator3>'].strip() != '':
        document.tables[0].add_row()
        cell = document.tables[0].rows[-1].cells[0]
        cell.merge(document.tables[0].rows[-2].cells[0])

        cell = document.tables[0].rows[-1].cells[1]
        
        cell.vertical_alignment = document.tables[0].rows[-2].cells[1].vertical_alignment
        base_paragraph = document.tables[0].rows[-2].cells[1].paragraphs[0]
        cell.paragraphs[0].alignment = base_paragraph.alignment

        run = cell.paragraphs[0].add_run()
        run.text = evaluation_info['<evaluator3>']

        base_font = document.tables[0].rows[-2].cells[1].paragraphs[0].runs[0].font
        run.font.name = base_font.name
        run.font.size = base_font.size
        run.font.bold = base_font.bold
        run.font.underline = base_font.underline
        run.font.color.rgb = base_font.color.rgb
    '''

    for row in document.tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for label in evaluation_info.keys():
                        if label in run.text:
                            run.text = run.text.replace(label, evaluation_info[label])
                            break

    for i, cell in enumerate(document.tables[1].rows[1].cells):
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.text = page_evaluations[0][i]
    
    for page_evaluation in page_evaluations[1:]:
        document.tables[1].add_row()
        for i, cell in enumerate(document.tables[1].rows[-1].cells):
            cell.vertical_alignment = document.tables[1].rows[1].cells[i].vertical_alignment
            base_paragraph = document.tables[1].rows[1].cells[i].paragraphs[0]
            cell.paragraphs[0].paragraph_format.alignment = base_paragraph.paragraph_format.alignment
            cell.paragraphs[0].paragraph_format.left_indent = base_paragraph.paragraph_format.left_indent
            cell.paragraphs[0].paragraph_format.right_indent = base_paragraph.paragraph_format.right_indent

            run = cell.paragraphs[0].add_run()
            run.text = page_evaluation[i]

            base_font = document.tables[1].rows[1].cells[i].paragraphs[0].runs[0].font
            run.font.name = base_font.name
            run.font.size = base_font.size
            run.font.bold = base_font.bold
            run.font.underline = base_font.underline
            run.font.color.rgb = base_font.color.rgb

    return document

if __name__ == '__main__' :
    make_list_of_pages_document()